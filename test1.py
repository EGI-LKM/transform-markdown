import os
import pandas as pd
import pdfplumber
from docx import Document


class MarkdownifyMCP:
    def convert(self, input_path, output_dir):
        """主转换方法"""
        ext = os.path.splitext(input_path)[1].lower()

        if ext == '.docx':
            return self._convert_docx(input_path, output_dir)
        elif ext == '.pdf':
            return self._convert_pdf(input_path, output_dir)
        elif ext in ('.xlsx', '.xls'):
            return self._convert_excel(input_path, output_dir)
        else:
            raise ValueError("Unsupported file format")

    def _convert_docx(self, docx_path, output_dir):
        """处理 Word 文档"""
        doc = Document(docx_path)
        md_content = []

        for para in doc.paragraphs:
            # 标题检测
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[1])
                md_content.append(f"{'#' * level} {para.text}")
            else:
                md_content.append(para.text)

        # 表格处理
        for table in doc.tables:
            md_table = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                md_table.append("| " + " | ".join(cells) + " |")

            # 添加表头分隔符
            if md_table:
                md_table.insert(1, "| " + " | ".join(["---"] * len(table.columns)) + " |")

            md_content.extend(["\n", "### 表格提取\n"] + md_table)

        self._save_md(md_content, docx_path, output_dir)

    def _convert_pdf(self, pdf_path, output_dir):
        """处理 PDF 文档"""
        md_content = ["## PDF 内容提取\n"]

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # 提取文本
                md_content.append(page.extract_text())

                # 提取表格
                for table in page.extract_tables():
                    df = pd.DataFrame(table[1:], columns=table[0])
                    md_content.append("\n### PDF表格\n" + df.to_markdown(index=False))

        self._save_md(md_content, pdf_path, output_dir)

    def _convert_excel(self, excel_path, output_dir):
        """处理 Excel 文件"""
        md_content = ["## Excel 表格转换\n"]

        xl = pd.ExcelFile(excel_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            md_content.append(f"### 工作表: {sheet_name}")
            md_content.append(df.to_markdown(index=False) + "\n")

        self._save_md(md_content, excel_path, output_dir)

    def _save_md(self, content, input_path, output_dir):
        """保存为 Markdown 文件"""
        base_name = os.path.basename(input_path)
        output_path = os.path.join(output_dir, f"{os.path.splitext(base_name)[0]}.md")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(content))

        print(f"转换完成 ➜ {output_path}")
